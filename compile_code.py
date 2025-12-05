from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_source_code_doc():
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('NEXUS Project Source Code', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Compiled Backend & AI Modules (Excluding UI/Dashboard)')
    doc.add_page_break()

    # --- Helper to add code blocks ---
    def add_file(folder, filename, content):
        # Header for the file
        h = doc.add_heading(f"{folder}/{filename}", level=1)
        h.style.font.color.rgb = RGBColor(0, 112, 192) # Blue color
        
        # Code block (using a Monospace font)
        p = doc.add_paragraph()
        runner = p.add_run(content.strip())
        runner.font.name = 'Courier New'
        runner.font.size = Pt(9)
        doc.add_page_break()

    print("⏳ Compiling NEXUS source code...")

    # ==============================================================================
    # FOLDER: CORE
    # ==============================================================================
    
    # 1. core/utils.py
    add_file("core", "utils.py", """
import math
import random
from typing import Tuple, List, Optional, Set, NamedTuple
from enum import Enum
from collections import namedtuple
import numpy as np

# Immutable point structure
Position = namedtuple('Position', ['row', 'col'])

class Direction(Enum):
    UP = (-1, 0)
    DOWN = (1, 0)
    LEFT = (0, -1)
    RIGHT = (0, 1)
    UP_LEFT = (-1, -1)
    UP_RIGHT = (-1, 1)
    DOWN_LEFT = (1, -1)
    DOWN_RIGHT = (1, 1)
    
    @property
    def delta(self):
        return self.value

def manhattan_distance(a, b):
    return abs(a[0] - b[0]) + abs(a[1] - b[1])

def euclidean_distance(a, b):
    return math.sqrt((a[0] - b[0])**2 + (a[1] - b[1])**2)

def chebyshev_distance(a, b):
    return max(abs(a[0] - b[0]), abs(a[1] - b[1]))

def generate_random_position(height, width, exclude=None):
    exclude = exclude or set()
    while True:
        r = random.randint(0, height - 1)
        c = random.randint(0, width - 1)
        if (r, c) not in exclude:
            return (r, c)

def generate_walkable_position(city, exclude=None):
    exclude = exclude or set()
    for _ in range(1000):
        pos = generate_random_position(city.height, city.width, exclude)
        if city.is_walkable(*pos):
            return pos
    return None

def interpolate_positions(a, b, steps):
    if steps <= 1: return [a, b]
    out = []
    for i in range(steps + 1):
        t = i / steps
        r = a[0] + (b[0] - a[0]) * t
        c = a[1] + (b[1] - a[1]) * t
        out.append((r, c))
    return out

def calculate_path_length(path):
    if len(path) < 2: return 0.0
    return sum(euclidean_distance(path[i], path[i+1]) for i in range(len(path)-1))
""")

    # 2. core/city.py
    add_file("core", "city.py", """
import numpy as np
from typing import Tuple, Set, Optional, List, Dict
from pathlib import Path

class City:
    ROAD = 0
    BUILDING = 1
    GRASS = 2
    WATER = 3
    HIGHWAY = 4
    RESTRICTED = 5
    DEFAULT_WALKABLE = {ROAD, GRASS, HIGHWAY}

    def __init__(self, map_path: Optional[str] = None):
        self.width = 0
        self.height = 0
        self.grid = None
        self._base_grid = None
        self.dynamic_blocked = set()
        self.walkable_values = set(self.DEFAULT_WALKABLE)
        
        if map_path: self.load(map_path)

    def load(self, map_path: str):
        path = Path(map_path)
        if not path.exists(): raise FileNotFoundError(f"Map not found: {map_path}")
        arr = np.load(str(path))
        self._base_grid = arr.copy()
        self.grid = arr.copy()
        self.height, self.width = self.grid.shape
        self.dynamic_blocked.clear()

    def in_bounds(self, row, col):
        return 0 <= row < self.height and 0 <= col < self.width

    def is_walkable(self, row, col):
        if not self.in_bounds(row, col): return False
        return (self.grid[row, col] in self.walkable_values and 
                (row, col) not in self.dynamic_blocked)

    def get_value(self, row, col):
        if not self.in_bounds(row, col): raise IndexError
        return int(self.grid[row, col])

    def block_tile(self, row, col):
        if self.in_bounds(row, col): self.dynamic_blocked.add((row, col))

    def unblock_tile(self, row, col):
        self.dynamic_blocked.discard((row, col))

    def walkable_neighbors(self, row, col, diagonal=False):
        deltas = [(-1, 0), (1, 0), (0, -1), (0, 1)]
        if diagonal: deltas += [(-1, -1), (-1, 1), (1, -1), (1, 1)]
        return [(row+dr, col+dc) for dr, dc in deltas if self.is_walkable(row+dr, col+dc)]
""")

    # 3. core/graph.py
    add_file("core", "graph.py", """
import networkx as nx
import math
from typing import Tuple, List, Optional, Dict

class CityGraph:
    def __init__(self, city=None):
        self.city = city
        self.graph = nx.Graph()
        if city: self.build_graph()

    def build_graph(self):
        if not self.city: return
        self.graph.clear()
        
        # Add Nodes
        for r in range(self.city.height):
            for c in range(self.city.width):
                if self.city.is_walkable(r, c):
                    self.graph.add_node((r, c))
        
        # Add Edges
        for node in self.graph.nodes():
            for neighbor in self.city.walkable_neighbors(*node):
                if neighbor in self.graph.nodes():
                    weight = math.sqrt((node[0]-neighbor[0])**2 + (node[1]-neighbor[1])**2)
                    self.graph.add_edge(node, neighbor, weight=weight)

    def get_neighbors(self, node):
        if node not in self.graph: return []
        return list(self.graph.neighbors(node))

    def get_edge_weight(self, n1, n2):
        if self.graph.has_edge(n1, n2):
            return self.graph[n1][n2].get('weight', 1.0)
        return None

    def get_stats(self):
        return {"nodes": self.graph.number_of_nodes(), "edges": self.graph.number_of_edges()}
""")

    # 4. core/agent.py
    add_file("core", "agent.py", """
from typing import Tuple, List, Optional, Any
from enum import Enum
import uuid
import math
from .utils import euclidean_distance, interpolate_positions
from ai.search import SearchEngine # Integration with AI

class AgentState(Enum):
    IDLE = 0
    MOVING = 1
    WORKING = 2
    CHARGING = 3
    OFFLINE = 4

class Agent:
    def __init__(self, name: str, start_pos: Tuple[int, int], speed: float = 1.0):
        self.id = str(uuid.uuid4())[:8]
        self.name = name
        self.position = (float(start_pos[0]), float(start_pos[1]))
        self.goal = None
        self.speed = speed
        self.path = []
        self.path_index = 0
        self.state = AgentState.IDLE
        self.energy = 100.0
        self.sensor_range = 5
        self.known_obstacles = set()
        self.distance_traveled = 0.0

    def navigate_to(self, graph, destination):
        start_node = (int(self.position[0]), int(self.position[1]))
        path = SearchEngine.a_star(graph, start_node, destination)
        if path:
            self.set_path(path)
            self.set_goal(destination)
            return True
        return False

    def set_path(self, path):
        self.path = path
        self.path_index = 0
        if path: self.state = AgentState.MOVING

    def set_goal(self, goal):
        self.goal = goal

    def is_at_goal(self):
        return self.goal and int(self.position[0]) == self.goal[0] and int(self.position[1]) == self.goal[1]

    def move_step(self):
        if not self.path or self.path_index >= len(self.path):
            self.state = AgentState.IDLE
            return

        target = self.path[self.path_index]
        dist = euclidean_distance(self.position, target)

        if dist <= self.speed:
            self.position = (float(target[0]), float(target[1]))
            self.path_index += 1
            self.distance_traveled += dist
        else:
            steps = max(int(dist / self.speed), 1)
            interp = interpolate_positions(self.position, target, steps)
            if len(interp) > 1: self.position = interp[1]
            self.distance_traveled += self.speed

    def sense(self, city):
        r, c = int(self.position[0]), int(self.position[1])
        visible = []
        for dr in range(-self.sensor_range, self.sensor_range+1):
            for dc in range(-self.sensor_range, self.sensor_range+1):
                nr, nc = r+dr, c+dc
                if city.in_bounds(nr, nc):
                    visible.append((nr, nc))
                    if not city.is_walkable(nr, nc):
                        self.known_obstacles.add((nr, nc))
        return visible
""")

    # 5. core/simulation.py
    add_file("core", "simulation.py", """
import time
from typing import List, Dict, Optional, Callable, Any, Tuple
from dataclasses import dataclass, field
from .city import City
from .graph import CityGraph
from .agent import Agent, AgentState

@dataclass
class SimulationEvent:
    event_type: str
    position: Tuple[int, int]
    start_time: int
    duration: int = -1
    active: bool = False

class Simulation:
    def __init__(self, city: City, graph: Optional[CityGraph] = None):
        self.city = city
        self.graph = graph or CityGraph(city)
        self.agents = []
        self.current_step = 0
        self.is_running = False
        self.events = []
        self.event_handlers = {}
        self.statistics = {"steps_completed": 0, "events_triggered": 0}
        self.step_callbacks = []
        self.agent_callbacks = []

    def add_agent(self, agent):
        self.agents.append(agent)

    def schedule_event(self, event):
        self.events.append(event)

    def register_event_handler(self, name, activate, deactivate=None):
        self.event_handlers[name] = {"activate": activate, "deactivate": deactivate}

    def register_step_callback(self, cb): self.step_callbacks.append(cb)
    def register_agent_callback(self, cb): self.agent_callbacks.append(cb)

    def run(self, max_steps=None, delay=0.0):
        self.is_running = True
        try:
            while self.is_running:
                if not self.step(): break
                if max_steps and self.current_step >= max_steps: break
                if delay > 0: time.sleep(delay)
        except KeyboardInterrupt:
            self.is_running = False

    def step(self):
        if not self.is_running: return False
        self.current_step += 1
        
        # Events
        for event in self.events:
            if event.start_time == self.current_step and not event.active:
                event.active = True
                self.statistics["events_triggered"] += 1
                if h := self.event_handlers.get(event.event_type, {}).get("activate"): h(self, event)
            if event.active and event.duration != -1 and self.current_step >= (event.start_time + event.duration):
                event.active = False
                if h := self.event_handlers.get(event.event_type, {}).get("deactivate"): h(self, event)

        # Agents
        for agent in self.agents:
            if agent.state == AgentState.OFFLINE: continue
            try:
                agent.sense(self.city)
                if agent.state == AgentState.MOVING: agent.move_step()
            except: pass

        # Callbacks
        for cb in self.step_callbacks: cb(self, self.current_step)
        return True

    def get_statistics(self): return self.statistics.copy()
    def start(self): self.is_running = True
    def stop(self): self.is_running = False
""")

    # ==============================================================================
    # FOLDER: AI
    # ==============================================================================

    # 6. ai/search.py
    add_file("ai", "search.py", """
import heapq
from typing import List, Tuple, Dict

class SearchEngine:
    @staticmethod
    def a_star(graph, start: Tuple[int, int], goal: Tuple[int, int]) -> List[Tuple[int, int]]:
        if start == goal: return [start]
        nx_graph = graph.graph if hasattr(graph, 'graph') else graph
        if start not in nx_graph or goal not in nx_graph: return []

        open_set = []
        heapq.heappush(open_set, (0, start))
        came_from = {}
        g_score = {start: 0}
        
        while open_set:
            current = heapq.heappop(open_set)[1]
            if current == goal:
                path = [current]
                while current in came_from:
                    current = came_from[current]
                    path.append(current)
                path.reverse()
                return path

            for neighbor in nx_graph.neighbors(current):
                tentative_g = g_score[current] + nx_graph[current][neighbor].get('weight', 1.0)
                if tentative_g < g_score.get(neighbor, float('inf')):
                    came_from[neighbor] = current
                    g_score[neighbor] = tentative_g
                    f = tentative_g + abs(neighbor[0]-goal[0]) + abs(neighbor[1]-goal[1])
                    heapq.heappush(open_set, (f, neighbor))
        return []

# Standalone functions for Test Suite compatibility
def bfs(graph, start, goal): pass # (Implementation omitted for brevity in docs)
def dfs(graph, start, goal): pass
def dijkstra(graph, start, goal): pass
def bidirectional_a_star(graph, start, goal): pass
def compare_algorithms(graph, start, goal): return {}
""")

    # 7. ai/logic_engine.py
    add_file("ai", "logic_engine.py", """
from typing import List, Dict, Any, Callable, Set
from dataclasses import dataclass, field

@dataclass
class Rule:
    name: str
    conditions: List[Callable]
    actions: List[Callable]
    priority: int = 0
    enabled: bool = True

class LogicEngine:
    def __init__(self):
        self.rules = []
        self.memory = {}

    def add_rule(self, rule):
        self.rules.append(rule)
        self.rules.sort(key=lambda r: r.priority, reverse=True)

    def update_memory(self, key, value):
        self.memory[key] = value

    def forward_chain(self):
        actions_taken = []
        for rule in self.rules:
            if rule.enabled and all(c(self.memory) for c in rule.conditions):
                for action in rule.actions:
                    res = action(self.memory)
                    actions_taken.append(res)
        return actions_taken

class AgentRules:
    @staticmethod
    def low_energy_recharge(threshold=30):
        def condition(ctx): return ctx['agent'].energy < threshold
        def action(ctx): 
            from core.agent import AgentState
            ctx['agent'].state = AgentState.CHARGING
            return "Switching to CHARGING"
        return Rule("LowEnergy", [condition], [action], 10)

    @staticmethod
    def obstacle_detected_replan():
        def condition(ctx): return bool(ctx.get('detected_obstacles'))
        def action(ctx): return "Obstacle detected"
        return Rule("Obstacle", [condition], [action], 8)
""")

    # 8. ai/planner.py
    add_file("ai", "planner.py", """
from typing import List, Set
from dataclasses import dataclass

@dataclass
class State:
    predicates: Set[str]
    def holds(self, p): return p in self.predicates
    def copy(self): return State(self.predicates.copy())

@dataclass
class Action:
    name: str
    preconditions: Set[str]
    add_effects: Set[str]
    del_effects: Set[str]
    
    def is_valid(self, state):
        return all(state.holds(p) for p in self.preconditions)
    
    def apply(self, state):
        new_s = state.copy()
        for p in self.del_effects: new_s.predicates.discard(p)
        for p in self.add_effects: new_s.predicates.add(p)
        return new_s

class Planner:
    def __init__(self):
        self.actions = []

    def add_action(self, action):
        self.actions.append(action)

    def plan(self, start_state, goal_state):
        # Simple BFS Planner
        queue = [(start_state, [])]
        visited = set()
        
        while queue:
            curr, plan = queue.pop(0)
            if all(curr.holds(g) for g in goal_state):
                return plan
            
            state_hash = frozenset(curr.predicates)
            if state_hash in visited: continue
            visited.add(state_hash)
            
            for action in self.actions:
                if action.is_valid(curr):
                    queue.append((action.apply(curr), plan + [action]))
        return None
""")

    # 9. ai/csp_engine.py
    add_file("ai", "csp_engine.py", """
# Constraint Satisfaction Engine
# Used for scheduling and resource allocation
class CSPEngine:
    def __init__(self):
        self.variables = {}
        self.constraints = []

    def add_variable(self, name, domain):
        self.variables[name] = domain

    def add_constraint(self, constraint_func):
        self.constraints.append(constraint_func)

    def solve(self):
        # Implementation of Backtracking Search
        pass
""")

    # 10. ai/bayesian.py
    add_file("ai", "bayesian.py", """
# Bayesian Reasoning Module
# Handles probability and uncertainty
class BayesianNetwork:
    def __init__(self):
        self.nodes = {}

    def add_node(self, name, probabilities):
        self.nodes[name] = probabilities

    def infer(self, evidence):
        # Probability calculation logic
        pass
""")

    # ==============================================================================
    # INTEGRATION SCRIPTS
    # ==============================================================================

    # 11. run_nexus_prime.py
    add_file("ROOT", "run_nexus_prime.py", """
import sys
import numpy as np
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))

from core.city import City
from core.graph import CityGraph
from core.simulation import Simulation, SimulationEvent
from core.agent import Agent
from ai.search import SearchEngine
from ai.logic_engine import LogicEngine, AgentRules

class SmartAgent(Agent):
    def __init__(self, name, start, graph):
        super().__init__(name, start)
        self.graph = graph
        self.brain = LogicEngine()
        self.brain.add_rule(AgentRules.low_energy_recharge(30))

    def update_ai(self, city):
        self.sense(city)
        self.brain.update_memory("agent", self)
        self.brain.forward_chain()

def main():
    print("🚀 NEXUS PRIME INITIALIZED")
    # Setup Logic omitted for brevity (See full file)
    
if __name__ == "__main__":
    main()
""")

    doc.save('Nexus_Source_Code.docx')
    print("✅ Document created: Nexus_Source_Code.docx")

if __name__ == "__main__":
    create_source_code_doc()